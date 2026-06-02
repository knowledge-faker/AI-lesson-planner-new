import os, json, asyncio
from docx import Document
from openai import AsyncOpenAI

from pydantic import BaseModel  # 🔥 必须导入这个

# --- 定义请求数据模型 ---
class GenerateReq(BaseModel):  # 🔥 加上这个类定义
    topic: str
    grade: str
    subject: str

_DEEPSEEK_KEY = os.getenv("DEEPSEEK_API_KEY", "").strip()
_DEEPSEEK_BASE = os.getenv("DEEPSEEK_BASE_URL", "https://api.deepseek.com").rstrip("/")
if not _DEEPSEEK_KEY:
    raise RuntimeError(
        "DEEPSEEK_API_KEY 未设置。在 .env 中配置 DeepSeek 控制台发放的 API Key。"
    )
client = AsyncOpenAI(api_key=_DEEPSEEK_KEY, base_url=_DEEPSEEK_BASE)

# 子任务1：生成教案数据
async def get_plan_data(topic, grade, subject):
    prompt = f"你是一名{grade}{subject}特级教师。请为《{topic}》编写一份详细教案，包含教学环节、内容、设计意图。返回JSON格式: {{'lesson_plan': [{{'环节':'', '内容':'', '设计意图':''}}]}}"
    response = await client.chat.completions.create(
        model="deepseek-chat",
        messages=[{"role": "user", "content": prompt}],
        response_format={"type": "json_object"}
    )
    return json.loads(response.choices[0].message.content)

# 子任务2：生成试卷数据
async def get_quiz_data(topic, grade, subject):
    prompt = f"你是一名{grade}{subject}命题专家。请为《{topic}》出一份100分制达标试卷，包含单选、填空、简答。必须标注分值且总分100。返回JSON格式: {{'quiz': [{{'题干':'', '选项':'', '答案':'', '分值':''}}]}}"
    response = await client.chat.completions.create(
        model="deepseek-chat",
        messages=[{"role": "user", "content": prompt}],
        response_format={"type": "json_object"}
    )
    return json.loads(response.choices[0].message.content)

# 子任务3：生成H5代码
async def get_html_data(topic, grade, subject):
    prompt = f"请为《{topic}》制作一个精美的H5交互课件代码，带内联CSS动画和JS交互逻辑。直接返回JSON格式: {{'html_content': '...html代码...'}}"
    response = await client.chat.completions.create(
        model="deepseek-chat",
        messages=[{"role": "user", "content": prompt}],
        response_format={"type": "json_object"}
    )
    return json.loads(response.choices[0].message.content)

async def generate_lesson_package(topic, grade, subject, user_id):
    # 🔥 并发执行三个任务，时间大大缩短
    plan_task, quiz_task, html_task = await asyncio.gather(
        get_plan_data(topic, grade, subject),
        get_quiz_data(topic, grade, subject),
        get_html_data(topic, grade, subject)
    )

    # --- 生成教案 Word ---
    doc_plan = Document()
    doc_plan.add_heading(f'{grade}{subject}教案：{topic}', 0)
    for p in plan_task.get('lesson_plan', []):
        doc_plan.add_heading(p.get('环节', '步骤'), level=1)
        doc_plan.add_paragraph(p.get('内容', ''))

    # --- 生成练习题 Word ---
    doc_quiz = Document()
    doc_quiz.add_heading(f'{grade}{subject}达标测试：{topic} (100分制)', 0)
    for q in quiz_task.get('quiz', []):
        p = doc_quiz.add_paragraph()
        p.add_run(f"题目：{q.get('题干', '')} ({q.get('分值', 0)}分)").bold = True
        doc_quiz.add_paragraph(f"答案：{q.get('答案', '')}")

    file_plan = f"教案_{grade}_{subject}_{topic}.docx"
    file_quiz = f"练习题_{grade}_{subject}_{topic}.docx"
    file_html = f"课件_{grade}_{subject}_{topic}.html"

    os.makedirs("files", exist_ok=True)
    doc_plan.save(f"files/{file_plan}")
    doc_quiz.save(f"files/{file_quiz}")
    with open(f"files/{file_html}", "w", encoding="utf-8") as f:
        f.write(html_task.get('html_content', ''))

    return {"plan": file_plan, "quiz": file_quiz, "html": file_html}